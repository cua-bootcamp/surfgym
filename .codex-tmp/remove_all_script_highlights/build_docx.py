from __future__ import annotations

import argparse
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


BODY_FONT = "Cambria"
BODY_SIZE = 12


def set_run_font(run, *, size=BODY_SIZE, bold=False, italic=False, underline=False, color=None, highlight=False):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)


def add_paragraph(doc, segments, *, align=None, before=0, after=8, line=1.05, keep=False, page_break=False):
    p = doc.add_paragraph()
    if page_break:
        br = p.add_run()
        br.add_break(WD_BREAK.PAGE)
        set_run_font(br)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.keep_together = keep
    for text, opts in segments:
        run = p.add_run(text)
        set_run_font(run, **opts)
    return p


def build_initial(output_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(BODY_SIZE)
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)

    doc.core_properties.title = "Community Workshop Reminder Call Guide"
    doc.core_properties.subject = "Workshop enrollment follow-up script"
    doc.core_properties.author = "Community Learning Office"
    doc.core_properties.keywords = "phone script, workshop, enrollment"

    add_paragraph(
        doc,
        [("Community Workshop Reminder Call Guide", {"size": 15, "bold": True})],
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=16,
        keep=True,
    )
    add_paragraph(
        doc,
        [
            ("This sample call guide supports follow-up with residents who asked about community technology classes. Use it as a guide only. ", {}),
            ("If you adapt it, review every bracketed note for your program.", {"bold": True, "italic": True}),
            (" Visit the staff portal for the current ", {}),
            ("enrollment checklist", {"color": "0000FF", "underline": True}),
            (".", {}),
        ],
        after=18,
    )
    add_paragraph(doc, [("Program Team:", {"bold": True}), (" Good morning, may I please speak with [", {}), ("participant name", {"italic": True}), ("]?", {})])
    add_paragraph(doc, [("If the person is not available: Thank whoever answered and ask for a better time to call.", {"italic": True})], before=5)
    add_paragraph(doc, [("If the person is available: Confirm that you have reached the correct person before continuing.", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" This is [", {}),
            ("caller name", {"italic": True}),
            ("] from ", {}),
            ("Harborview Learning Center", {"highlight": True}),
            (" [", {}),
            ("program team", {"italic": True}),
            ("]. I coordinate the ", {}),
            ("Digital Basics", {"highlight": True}),
            (" workshop with [", {}),
            ("instructor name", {"italic": True}),
            ("]. Is now a convenient time to speak for a few minutes?", {}),
        ],
    )
    add_paragraph(doc, [("If the person says “No” or “I’m busy”", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Of course. [", {}),
            ("Offer two callback windows. If the person sounds unsure, thank them and end the call.", {"italic": True}),
            ("]", {}),
        ],
    )
    add_paragraph(doc, [("If the person says “Yes”", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Great. You recently asked about help with everyday computer skills at ", {}),
            ("Harborview Learning Center", {"highlight": True}),
            (" [", {}),
            ("branch name", {"italic": True}),
            ("]. The ", {}),
            ("Digital Basics", {"highlight": True}),
            (" series covers secure passwords, email attachments, online forms, and video appointments. Would you like a short overview?", {}),
        ],
    )
    add_paragraph(doc, [("If the person says “No” or “Not right now”", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" That is completely fine. We only wanted to follow up on your request from [", {}),
            ("date", {"italic": True}),
            ("]. If you change your mind, the ", {}),
            ("Harborview Learning Center", {"highlight": True}),
            (" information desk can explain future sessions. Thank you for your time.", {}),
        ],
    )
    add_paragraph(doc, [("If the person would like the overview", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" The next ", {}),
            ("Digital Basics", {"highlight": True}),
            (" series meets on four Tuesday afternoons. Each session is held in the second-floor ", {}),
            ("laptop lab", {"highlight": True}),
            (" and lasts about ninety minutes. The group is limited to twelve participants so instructors can provide individual help.", {}),
        ],
    )
    add_paragraph(
        doc,
        [
            ("Participants practise on centre computers and may bring a personal phone, tablet, or laptop for selected activities. The ", {}),
            ("laptop lab", {"highlight": True}),
            (" has adjustable desks, large-print keyboards, headphones, and a screen reader station. No previous class experience is required, and every activity can be repeated at home.", {}),
        ],
    )
    add_paragraph(
        doc,
        [
            ("The first meeting begins with a short skills survey, a review of shared-computer privacy, and time to check each participant’s device. Staff at ", {}),
            ("Harborview Learning Center", {"highlight": True}),
            (" can arrange a loaner device when one is requested before [", {}),
            ("request deadline", {"italic": True}),
            ("].", {}),
        ],
        page_break=True,
    )
    add_paragraph(
        doc,
        [
            ("Registration includes all four sessions, printed practice sheets, and two optional drop-in clinics. The clinics take place in the same ", {}),
            ("laptop lab", {"highlight": True}),
            (" and are intended for questions that come up between classes. Participants are welcome to bring a support person to the first session.", {}),
        ],
    )
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Before I reserve a place, I need to confirm your preferred name, a telephone number, and whether you would like any accessibility support. You may skip any optional question and you can ask us to remove your registration at any time.", {}),
        ],
    )
    add_paragraph(doc, [("If the person is ready to register", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" May I confirm the spelling of your name? [", {}),
            ("Record the response", {"italic": True}),
            (".] Which telephone number should we use for a weather closure or schedule change? [", {}),
            ("Record the response", {"italic": True}),
            (".]", {}),
        ],
    )
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Would large-print materials, step-free access, hearing support, or another adjustment make the workshop easier to attend? [", {}),
            ("Record only what the person chooses to share and pass the request to the access coordinator.", {"italic": True}),
            ("]", {}),
        ],
    )
    add_paragraph(doc, [("If the person needs time before registering", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" I can email or post the class summary for you to review. The summary lists session dates, transport information, accessibility contacts, and what to bring. I will hold a provisional place until [", {}),
            ("hold date", {"italic": True}),
            ("], after which the place returns to the waiting list.", {}),
        ],
    )
    add_paragraph(doc, [("If the person asks for an email copy", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" What email address should I use? ____________________ Thank you. I will send the summary from workshops@harborview-learning.example. Please check your junk folder if it does not arrive by tomorrow afternoon.", {}),
        ],
    )
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Do you have any questions about the schedule, the activities, or getting to ", {}),
            ("Harborview Learning Center", {"highlight": True}),
            ("?", {}),
        ],
    )
    add_paragraph(doc, [("Answer the person’s questions using the current program notes. If you are unsure, record the question and offer a callback.", {"italic": True})], before=5)
    add_paragraph(
        doc,
        [
            ("Program Team:", {"bold": True}),
            (" Thank you for your time. We will send the information you requested and contact you again only if you asked us to. Have a good day.", {}),
        ],
        after=0,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def clear_highlights(initial_path: Path, gold_path: Path) -> int:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    gold_path.parent.mkdir(parents=True, exist_ok=True)
    removed = 0
    with zipfile.ZipFile(initial_path, "r") as src, zipfile.ZipFile(gold_path, "w", zipfile.ZIP_DEFLATED) as dst:
        for info in src.infolist():
            data = src.read(info.filename)
            if info.filename == "word/document.xml":
                root = etree.fromstring(data)
                highlights = root.xpath(".//w:highlight", namespaces=ns)
                removed = len(highlights)
                for node in highlights:
                    node.getparent().remove(node)
                data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
            dst.writestr(info, data)
    return removed


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--initial", type=Path, required=True)
    parser.add_argument("--gold", type=Path, required=True)
    args = parser.parse_args()

    build_initial(args.initial)
    count = clear_highlights(args.initial, args.gold)
    if count != 9:
        raise RuntimeError(f"Expected 9 highlighted spans, found {count}")
    print(f"Created {args.initial}")
    print(f"Created {args.gold}")
    print(f"Removed {count} highlight nodes from Gold")


if __name__ == "__main__":
    main()
